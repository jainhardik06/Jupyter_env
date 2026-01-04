# First, you may need to install the library: pip install python-docx

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_hackathon_proposal():
    doc = Document()

    # --- Title Page Styling ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('Aegis-X: Decentralized Threat Intelligence', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('via Federated Graph Learning & Blockchain Provenance')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = 'Subtitle'

    doc.add_paragraph() # Spacer

    # Team Info
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Track 4: Inter-Agency Data Sharing\n').bold = True
    info.add_run('Team Name: [INSERT TEAM NAME]\n')
    info.add_run('Date: December 19, 2025')

    doc.add_page_break()

    # --- Section 1: Executive Abstract ---
    doc.add_heading('1. Executive Abstract', level=1)
    
    abstract_text = (
        "The primary failure mode in inter-agency counter-terrorism is data isolation; agencies possess "
        "complementary intelligence fragments but are operationally restricted from pooling sensitive raw data. "
        "Aegis-X solves this 'Privacy-Utility Paradox' by introducing a Federated Learning (FL) architecture "
        "that enables collaborative threat detection without ever moving raw data from an agency’s local server.\n\n"
        "Our solution operates on a three-stage modular pipeline. First, a Polyglot Ingestion Engine "
        "(built on Python & Apache Kafka) normalizes heterogeneous data streams. Second, a distributed "
        "Knowledge Graph utilizes Graph Neural Networks (GNN) to autonomously identify high-dimensional "
        "relationships—such as correlating a flagged financial transaction with a CCTV vehicle sighting—that "
        "human analysts would miss. Finally, a Zero-Trust Security Layer backed by a Blockchain Ledger "
        "ensures an immutable audit trail for every data interaction."
    )
    p = doc.add_paragraph(abstract_text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- Section 2: Problem Statement ---
    doc.add_heading('2. Problem Statement: The Silo Effect', level=1)
    problem_text = (
        "Current national security infrastructure suffers from the 'Blindness of Asymmetry'. "
        "The Police, Financial Intelligence Unit (FIU), and Border Control operate on disconnected databases. "
        "Manual data sharing requests take days to process, by which time critical threats may have escalated. "
        "Furthermore, agencies are hesitant to share full databases due to privacy laws and fear of leaks."
    )
    doc.add_paragraph(problem_text)

    # --- Section 3: System Architecture (Visual 1) ---
    doc.add_heading('3. System Architecture', level=1)
    
    # Placeholder for Chart 1
    p = doc.add_paragraph()
    run = p.add_run('[INSERT FLOWCHART 1 HERE: High-Level Modular Pipeline]')
    run.bold = True
    run.font.color.rgb = None # Default color
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        "The Aegis-X architecture is composed of three distinct layers:\n"
        "1. Ingestion Layer: Connects to agency silos via read-only APIs.\n"
        "2. Intelligence Layer: The GNN Engine that connects entities (Person, Vehicle, Phone) across silos.\n"
        "3. Governance Layer: The Blockchain network that logs every 'Grant Access' request."
    )

    # --- Section 4: Federated Learning Methodology (Visual 2) ---
    doc.add_heading('4. AI Methodology: Federated Learning', level=1)

    # Placeholder for Chart 2
    p = doc.add_paragraph()
    run = p.add_run('[INSERT FLOWCHART 2 HERE: The Federated Training Loop]')
    run.bold = True
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "To preserve privacy, we utilize a 'Compute-to-Data' approach. Instead of sending data to the cloud, "
        "we send the AI model to the data. Agency A and Agency B train the model locally on their private servers. "
        "Only the mathematical weight updates (gradients) are sent to the central server, ensuring raw records "
        "never leave the agency firewall."
    )

    # --- Section 5: Technical Stack ---
    doc.add_heading('5. Technology Stack', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology Used'
    
    stack = [
        ('Core AI Engine', 'TensorFlow Federated (TFF), PyTorch Geometric (GNN)'),
        ('Natural Language Processing', 'HuggingFace Transformers (BERT/RoBERTa)'),
        ('Knowledge Graph Database', 'Neo4j Community Edition'),
        ('Search Engine', 'Elasticsearch (Vector Search)'),
        ('Backend API', 'Python FastAPI (Async), Node.js'),
        ('Data Stream', 'Apache Kafka / RabbitMQ'),
        ('Frontend Dashboard', 'React.js + WebGL (react-force-graph)'),
        ('Security Audit', 'Hyperledger Fabric / Custom SHA-256 Blockchain')
    ]
    
    for component, tech in stack:
        row_cells = table.add_row().cells
        row_cells[0].text = component
        row_cells[1].text = tech

    doc.save('Aegis_X_Project_Proposal.docx')
    print("Document generated successfully!")

if __name__ == "__main__":
    create_hackathon_proposal()